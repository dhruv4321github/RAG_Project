"""
document_processor.py — Document Ingestion Pipeline

Handles the full lifecycle of turning an uploaded file into searchable vectors:
  1. Parse: Extract raw text from PDF, DOCX, or TXT files
  2. Chunk: Split text into overlapping segments using LangChain
  3. Embed: Generate vector embeddings via OpenAI
  4. Store: Save chunks and vectors to PostgreSQL/pgvector

This is called whenever a user uploads a new document via the API.
"""

import os
import uuid
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI
from sqlalchemy.orm import Session

from app.config import settings
from app.models.database import Document, DocumentChunk

# Initialize the OpenAI client for embedding generation
client = OpenAI(api_key=settings.OPENAI_API_KEY)


# ──────────────────────────────────────────────
# Step 1: Parse Documents
# ──────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> list[dict]:
    """
    Extracts text from a PDF file, page by page.
    Returns a list of dicts with 'text' and 'page_number' keys
    so we can track which page each chunk came from.
    """
    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append({"text": text.strip(), "page_number": i + 1})
    return pages


def extract_text_from_docx(file_path: str) -> list[dict]:
    """
    Extracts text from a DOCX file.
    DOCX files don't have clear page boundaries, so we treat
    the entire document as page 1.
    """
    doc = DocxDocument(file_path)
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [{"text": full_text, "page_number": 1}] if full_text else []


def extract_text_from_txt(file_path: str) -> list[dict]:
    """Reads a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read().strip()
    return [{"text": text, "page_number": 1}] if text else []


def extract_text(file_path: str, file_type: str) -> list[dict]:
    """
    Routes to the correct parser based on file type.
    Returns a list of {text, page_number} dicts.
    """
    parsers = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "txt": extract_text_from_txt,
    }
    parser = parsers.get(file_type)
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")
    return parser(file_path)


# ──────────────────────────────────────────────
# Step 2: Chunk Text
# ──────────────────────────────────────────────

def chunk_text(pages: list[dict]) -> list[dict]:
    """
    Splits extracted text into overlapping chunks using LangChain's
    RecursiveCharacterTextSplitter.

    Why overlapping chunks?
    If a key piece of information spans a chunk boundary (e.g., a sentence
    split between chunk 5 and chunk 6), the overlap ensures both chunks
    contain the full sentence. This improves retrieval accuracy.

    Returns a list of dicts: {content, page_number, chunk_index}
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]  # Tries to split at natural boundaries
    )

    chunks = []
    chunk_index = 0

    for page in pages:
        page_chunks = splitter.split_text(page["text"])
        for chunk_text in page_chunks:
            chunks.append({
                "content": chunk_text,
                "page_number": page["page_number"],
                "chunk_index": chunk_index,
            })
            chunk_index += 1

    return chunks


# ──────────────────────────────────────────────
# Step 3: Generate Embeddings
# ──────────────────────────────────────────────

def generate_embeddings(texts: list[str]) -> list[list[float]]:
    """
    Sends text chunks to OpenAI's embedding API in a batch.
    Returns a list of 1536-dimensional vectors (one per chunk).

    The embedding model converts text into a numerical representation
    where semantically similar texts have vectors that are close together.
    This is what enables similarity search in the RAG pipeline.
    """
    response = client.embeddings.create(
        model=settings.EMBEDDING_MODEL,
        input=texts,
    )
    # Sort by index to maintain order (API may return out of order)
    sorted_embeddings = sorted(response.data, key=lambda x: x.index)
    return [item.embedding for item in sorted_embeddings]


# ──────────────────────────────────────────────
# Step 4: Full Processing Pipeline
# ──────────────────────────────────────────────

def process_document(
    db: Session,
    file_path: str,
    filename: str,
    file_type: str,
    file_size: int,
) -> Document:
    """
    Orchestrates the full document ingestion pipeline:
      1. Create a Document record (status = "processing")
      2. Extract text from the file
      3. Split into chunks
      4. Generate embeddings for all chunks
      5. Store chunks + embeddings in the database
      6. Update document status to "ready"

    If anything fails, the document status is set to "error".
    """
    # Create the document record
    doc = Document(
        id=uuid.uuid4(),
        filename=filename,
        file_type=file_type,
        file_size_bytes=file_size,
        status="processing",
    )
    db.add(doc)
    db.commit()

    try:
        # Extract text
        pages = extract_text(file_path, file_type)
        if not pages:
            raise ValueError("No text could be extracted from the document")

        # Chunk the text
        chunks = chunk_text(pages)
        if not chunks:
            raise ValueError("No chunks were generated from the document")

        # Generate embeddings (batch for efficiency)
        chunk_texts = [c["content"] for c in chunks]

        # Process in batches of 100 to avoid API limits
        all_embeddings = []
        batch_size = 100
        for i in range(0, len(chunk_texts), batch_size):
            batch = chunk_texts[i:i + batch_size]
            batch_embeddings = generate_embeddings(batch)
            all_embeddings.extend(batch_embeddings)

        # Store chunks with their embeddings
        for chunk_data, embedding in zip(chunks, all_embeddings):
            db_chunk = DocumentChunk(
                id=uuid.uuid4(),
                document_id=doc.id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                page_number=chunk_data["page_number"],
                embedding=embedding,
            )
            db.add(db_chunk)

        # Update document status
        doc.chunk_count = len(chunks)
        doc.status = "ready"
        db.commit()

        return doc

    except Exception as e:
        # Mark document as errored so the UI can show the failure
        doc.status = "error"
        db.commit()
        raise e
    finally:
        # Clean up the uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)
